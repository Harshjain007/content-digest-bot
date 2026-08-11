"""Content extraction layer.

Each extractor returns a dict:
    {
        "source": "YouTube" | "Article" | "Instagram" | "Topic",
        "title":   str,
        "url":     str | None,
        "text":    str | None,   # the raw material Claude will digest
        # optional flags:
        "needs_caption": True,   # Instagram: ask user to paste caption
        "failed":        True,   # Article: extraction failed
        "user_note":     str,    # extra text the user typed alongside the link
    }
"""
import logging
import re

logger = logging.getLogger(__name__)

URL_RE = re.compile(r"https?://[^\s]+", re.IGNORECASE)


def classify(text):
    """Return (kind, url_or_None, leftover_text)."""
    urls = URL_RE.findall(text.strip())
    if not urls:
        return ("topic", None, text.strip())
    url = urls[0]
    rest = text.replace(url, "").strip()
    if "youtu.be" in url or "youtube.com" in url:
        return ("youtube", url, rest)
    if "instagram.com" in url or "instagr.am" in url:
        return ("instagram", url, rest)
    return ("article", url, rest)


# ----------------------------------------------------------------- YouTube
def _youtube_id(url):
    m = re.search(r"(?:v=|youtu\.be/|shorts/|embed/)([\w-]{11})", url)
    if m:
        return m.group(1)
    m = re.search(r"([\w-]{11})", url)
    return m.group(1) if m else None


def _yt_metadata(url):
    try:
        import yt_dlp
        with yt_dlp.YoutubeDL({"quiet": True, "skip_download": True,
                               "no_warnings": True}) as ydl:
            info = ydl.extract_info(url, download=False)
            return info.get("title"), info.get("description")
    except Exception as e:  # noqa: BLE001
        logger.warning("yt-dlp metadata failed: %s", e)
        return None, None


def _yt_transcript(vid):
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        api = YouTubeTranscriptApi()
        try:
            fetched = api.fetch(vid, languages=["en", "en-US", "en-GB"])
        except TypeError:
            fetched = api.fetch(vid)
        parts = [s.text for s in fetched]
        return " ".join(parts)
    except Exception as e:  # noqa: BLE001
        logger.warning("Transcript failed for %s: %s", vid, e)
        return None


def extract_youtube(url, max_chars=30000):
    vid = _youtube_id(url)
    title, description = _yt_metadata(url)
    transcript = _yt_transcript(vid) if vid else None
    text = (description or "") + "\n\n" + (transcript or "")
    text = text[:max_chars]
    return {
        "source": "YouTube",
        "title": title or (f"Video {vid}" if vid else url),
        "url": url,
        "text": text,
        "had_transcript": bool(transcript),
    }


# ----------------------------------------------------------------- Article
def _is_pdf(url, content_type=None):
    if url and url.lower().split("?")[0].endswith(".pdf"):
        return True
    return bool(content_type) and "application/pdf" in content_type.lower()


def _fetch_bytes(url, timeout=30):
    """Return (content_bytes, content_type) or (None, None)."""
    import requests
    headers = {"User-Agent": "Mozilla/5.0 (compatible; content-digest-bot/1.0)"}
    try:
        r = requests.get(url, headers=headers, timeout=timeout,
                         allow_redirects=True)
        r.raise_for_status()
        return r.content, r.headers.get("Content-Type", "")
    except Exception as e:  # noqa: BLE001
        logger.warning("PDF fetch failed for %s: %s", url, e)
        return None, None


def extract_pdf(url, max_chars=30000):
    """Extract text from a PDF URL via pypdf."""
    raw, ctype = _fetch_bytes(url)
    if not raw:
        return {"source": "Article", "title": "PDF link", "url": url,
                "text": None, "failed": True}
    try:
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(raw))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue
        text = "\n\n".join(p for p in parts if p).strip()
        title = (reader.metadata.title if reader.metadata
                 else None) or url
        if len(text) < 100:
            raise RuntimeError("Too little text extracted from PDF")
        return {
            "source": "Article",
            "title": title,
            "url": url,
            "text": text[:max_chars],
            "is_pdf": True,
        }
    except Exception as e:  # noqa: BLE001
        logger.warning("PDF extraction failed for %s: %s", url, e)
        return {"source": "Article", "title": "PDF link", "url": url,
                "text": None, "failed": True}


def _is_doc(url, content_type=None):
    u = (url or "").lower().split("?")[0]
    if u.endswith(".docx"):
        return True
    if u.endswith(".doc"):
        return True
    ct = (content_type or "").lower()
    return "wordprocessingml" in ct or "application/msword" in ct or \
        "application/vnd.openxmlformats-officedocument.wordprocessingml" in ct


def extract_doc(url, max_chars=30000):
    """Extract text from a Word document (.docx) URL via python-docx.

    .docx is fully supported. Legacy .doc (binary) is not readable by
    python-docx — we flag it so the bot can tell the user.
    """
    if url and url.lower().split("?")[0].endswith(".doc") and \
            not url.lower().endswith(".docx"):
        return {"source": "Article", "title": "Word doc (.doc)",
                "url": url, "text": None,
                "failed": True,
                "note": "Legacy .doc files aren't supported — please "
                        "re-save as .docx and reshare."}
    raw, ctype = _fetch_bytes(url)
    if not raw:
        return {"source": "Article", "title": "Word doc", "url": url,
                "text": None, "failed": True}
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull tables (common in docs).
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts).strip()
        title = (doc.core_properties.title or "").strip() or url
        if len(text) < 50:
            raise RuntimeError("Too little text extracted from document")
        return {
            "source": "Article",
            "title": title,
            "url": url,
            "text": text[:max_chars],
            "is_doc": True,
        }
    except Exception as e:  # noqa: BLE001
        logger.warning("DOC extraction failed for %s: %s", url, e)
        return {"source": "Article", "title": "Word doc", "url": url,
                "text": None, "failed": True}


def extract_article(url, max_chars=30000):
    # PDFs and PDF-like responses go through the PDF extractor.
    raw_head, ctype = _fetch_bytes(url)
    if raw_head is not None and _is_doc(url, ctype):
        return extract_doc(url, max_chars=max_chars)
    if raw_head is not None and _is_pdf(url, ctype):
        return extract_pdf(url, max_chars=max_chars)
    try:
        import trafilatura
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            raise RuntimeError("Could not fetch URL")
        text = trafilatura.extract(downloaded, include_comments=False,
                                   include_tables=True)
        meta = trafilatura.extract_metadata(downloaded)
        title = meta.title if meta else None
        if not text or len(text) < 200:
            raise RuntimeError("No main text extracted")
        return {
            "source": "Article",
            "title": title or url,
            "url": url,
            "text": text[:max_chars],
        }
    except Exception as e:  # noqa: BLE001
        logger.warning("Article extraction failed for %s: %s", url, e)
        return {"source": "Article", "title": "Article link", "url": url,
                "text": None, "failed": True}


# --------------------------------------------------------------- Instagram
def _ig_shortcode(url):
    m = re.search(r"instagram\.com/(?:reel|p|tv)/([\w-]+)", url)
    return m.group(1) if m else None


def extract_instagram(url, username=None, password=None, max_chars=30000):
    try:
        import instaloader
        shortcode = _ig_shortcode(url)
        if not shortcode:
            raise RuntimeError("Could not parse Instagram shortcode")
        L = instaloader.Instaloader()
        if username and password:
            try:
                L.login(username, password)
            except Exception as e:  # noqa: BLE001
                logger.warning("IG login failed: %s", e)
        post = instaloader.Post.from_shortcode(L, shortcode)
        caption = post.caption or ""
        text = (f"Caption:\n{caption}\n\n"
                f"(Instagram post by @{post.owner_username or 'unknown'}; "
                f"likes ~{post.likes_count})")
        return {
            "source": "Instagram",
            "title": (caption.split("\n")[0][:80] if caption
                      else "Instagram reel/post"),
            "url": url,
            "text": text[:max_chars],
        }
    except Exception as e:  # noqa: BLE001
        logger.warning("Instagram extraction failed for %s: %s", url, e)
        return {"source": "Instagram", "title": "Instagram link", "url": url,
                "text": None, "needs_caption": True}


# -------------------------------------------------------------------- Topic
def extract_topic(topic, max_chars=2000):
    return {"source": "Topic", "title": topic[:80], "url": None,
            "text": topic[:max_chars]}


# ----------------------------------------------------------------- dispatch
def extract(text):
    """Classify a raw message and return the appropriate content dict."""
    from .config import INSTAGRAM_USERNAME, INSTAGRAM_PASSWORD  # local import
    kind, url, rest = classify(text)

    if kind == "topic":
        data = extract_topic(text)
    elif kind == "youtube":
        data = extract_youtube(url)
    elif kind == "instagram":
        data = extract_instagram(url, INSTAGRAM_USERNAME, INSTAGRAM_PASSWORD)
    else:  # article
        data = extract_article(url)

    if rest:
        data["user_note"] = rest
    return data
