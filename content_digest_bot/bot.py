"""Telegram knowledge-keeper bot (long-polling — runs locally).

When you share a link/topic, the bot classifies it and:
  • TOOL  (article that links to GitHub, or a GitHub URL directly)
        -> summarize article + fetch GitHub README -> structured JSON:
           {title, description, problem, how this works, links}
  • LEARNING (upskilling / productivity / self-improvement article)
        -> structured JSON: {description, links, takeAways}
  • CONCEPT (a topic or other AI content)
        -> short summary, then "explain more?" -> full deep-dive (chat only)

All TOOL/LEARNING entries are de-duplicated against existing JSON and saved
to data/ (resources.json / learnings.json), mirrored into data.json for the
HTML viewer in site/.
"""
import logging
import os
import re
from datetime import datetime, timezone

from telegram import Update
from telegram.constants import ParseMode
from telegram.ext import (Application, CommandHandler, ContextTypes,
                          MessageHandler, filters)
from telegram import Document

from .config import TELEGRAM_BOT_TOKEN, ANTHROPIC_API_KEY, ANTHROPIC_MODEL
from .extractors import extract, classify, URL_RE
from .github_api import is_github_url, fetch_repo
from .synthesize import synthesize, synthesize_json
from .format_telegram import md_to_telegram_html, split_html
from .moderate import is_allowed, REJECT_MSG
from .store import add_resource, add_learning
from .prompts import build_tool_json_prompt, build_learning_json_prompt

# Private bot: only these Telegram user/chat IDs may use it.
# Your chat id was observed earlier; add others as needed.
ALLOWED_CHAT_IDS = {811501439}

# Live knowledge register (GitHub Pages) — shown after every save so you can
# open it from anywhere.
PAGES_URL = "https://harshjain007.github.io/content-digest-bot/site/index.html"

NOT_AUTHORIZED = ("🔒 This bot is private. You are not authorized to use it.")

logging.basicConfig(
    format="%(asctime)s %(levelname)s %(name)s: %(message)s",
    level=logging.INFO,
)
logger = logging.getLogger(__name__)

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
NOTES_DIR = os.path.join(REPO_ROOT, "notes")
os.makedirs(NOTES_DIR, exist_ok=True)

HELP = (
    "I'm your AI knowledge-keeper. Share:\n"
    "• an article that links to a GitHub repo (or a GitHub URL) → I save a "
    "structured tool card (what it is, problem, how to use, links)\n"
    "• an upskilling / productivity article → I save a learning card "
    "(what it solves, takeaways)\n"
    "• an AI concept / topic → I give a short summary; reply 'yes' to go deeper\n\n"
    "Everything is de-duplicated and saved locally as JSON for a web viewer."
)

EXPAND_WORDS = ("yes", "y", "more", "explain", "explain more", "details",
                "detail", "elaborate", "tell me more", "full", "deep", "go on")


def _is_expand(text):
    return any(w in text.strip().lower() for w in EXPAND_WORDS)


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_chat.id not in ALLOWED_CHAT_IDS:
        await update.message.reply_text(NOT_AUTHORIZED)
        return
    await update.message.reply_text("👋 " + HELP)


async def help_cmd(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.effective_chat.id not in ALLOWED_CHAT_IDS:
        await update.message.reply_text(NOT_AUTHORIZED)
        return
    await update.message.reply_text(HELP)


async def _send_html(update, text, prefix=""):
    for i, chunk in enumerate(split_html(md_to_telegram_html(text))):
        p = (f"<b>{prefix}</b>\n\n" if prefix and i == 0 else "")
        try:
            await update.message.reply_text(p + chunk, parse_mode=ParseMode.HTML)
        except Exception as e:  # noqa: BLE001
            logger.warning("HTML send failed: %s", e)
            await update.message.reply_text(p + re.sub(r"<[^>]+>", "", chunk))


def _pretty_json(obj):
    import json
    return "```json\n" + json.dumps(obj, indent=2, ensure_ascii=False) + "\n```"


# ----------------------------------------------------------------- handlers
def _normalize(text):
    """Strip the @url: wrapper, backticks and markdown fences users paste."""
    t = text.strip()
    # remove leading @url: / url: markers (case-insensitive)
    m = re.match(r"@?url:\s*", t, re.IGNORECASE)
    if m:
        t = t[m.end():]
    # strip code fences / backticks
    t = t.replace("```", "").replace("`", "")
    return t.strip()


async def handle(update: Update, context: ContextTypes.DEFAULT_TYPE):
    raw_text = update.message.text
    text = _normalize(raw_text)
    chat_id = update.effective_chat.id
    logger.info("Incoming from chat_id=%s: %s", chat_id, raw_text[:80])

    # Private bot guard.
    if chat_id not in ALLOWED_CHAT_IDS:
        logger.warning("Unauthorized chat %s blocked", chat_id)
        await update.message.reply_text(NOT_AUTHORIZED)
        return

    await context.bot.send_chat_action(chat_id, "typing")

    # Explain-more follow-up for a pending CONCEPT.
    pending = context.user_data.get("pending")
    if pending and _is_expand(text):
        try:
            full = synthesize(pending, user_note=pending.get("user_note"),
                              mode="full")
            if full is None:
                await update.message.reply_text("⚠️ Nothing to explain.")
                return
            await _send_html(update, full)
            context.user_data.pop("pending", None)
        except Exception as e:  # noqa: BLE001
            logger.exception("expand error")
            await update.message.reply_text(f"❌ Something went wrong: {e}")
        return

    # Topic gate only applies to bare topics (no link). Any shared URL —
    # GitHub repo, article, PDF, YouTube — is an explicit "file this" intent
    # and always passes through.
    has_url = bool(URL_RE.search(text))
    if not has_url and not is_allowed(text):
        logger.info("Rejected off-topic: %s", text[:60])
        await update.message.reply_text(REJECT_MSG, parse_mode=ParseMode.HTML)
        return

    status = await update.message.reply_text("🔍 Working…")
    try:
        kind, url, rest = classify(text)

        # ---- TOOL flow: GitHub URL, or article -> GitHub ----
        if is_github_url(url):
            await _handle_tool(update, status, github_url=url,
                               article_url=None, article_text=None)
            return
        if kind == "article":
            data = extract(text)
            if data.get("failed"):
                await status.edit_text("⚠️ Couldn't read that article. "
                                       "Paste the text and I'll digest it.")
                return
            # PDFs / Word docs: full descriptive summary, saved as a resource card.
            if data.get("is_pdf") or data.get("is_doc"):
                await _handle_pdf(update, status, data)
                return
            gh = _find_github_link(data.get("text") or "")
            if gh:
                await _handle_tool(update, status, github_url=gh,
                                   article_url=url,
                                   article_text=data.get("text"))
            else:
                # No GitHub link -> is it a learning article?
                if _looks_like_learning(data.get("text") or ""):
                    await _handle_learning(update, status,
                                           article_text=data.get("text"),
                                           article_url=url)
                else:
                    # Generic AI article: just summarize in chat.
                    await status.delete()
                    await _send_html(update,
                                    synthesize(data, mode="summary"),
                                    prefix="Want me to explain more? Reply 'yes'.")
                    context.user_data["pending"] = data
            return

        # ---- CONCEPT / topic / youtube / instagram ----
        data = extract(text)
        if data.get("needs_caption"):
            await status.edit_text("⚠️ Couldn't read this Instagram reel. "
                                   "Paste the caption and I'll digest it.")
            return
        if data.get("failed"):
            await status.edit_text("⚠️ Couldn't read that. Paste the text.")
            return
        context.user_data["pending"] = data
        await status.delete()
        await _send_html(update, synthesize(data, mode="summary"),
                        prefix="Want me to explain more? Reply 'yes'.")
    except Exception as e:  # noqa: BLE001
        logger.exception("handle error")
        await status.edit_text(f"❌ Something went wrong: {e}")


def _pages_link():
    return f"🔗 Register: {PAGES_URL}"


async def _handle_tool(update, status, github_url, article_url, article_text):
    repo = fetch_repo(github_url)
    gh_text = repo.get("readme") if repo else ""
    website = repo.get("homepage") if repo else ""
    if repo:
        github_url = repo.get("html_url") or github_url
    prompt = build_tool_json_prompt(article_text, gh_text, article_url,
                                    github_url, website)
    try:
        entry = synthesize_json(prompt, num_predict=1200)
    except Exception as e:  # noqa: BLE001
        logger.exception("tool json failed")
        await status.edit_text(f"❌ Couldn't build the tool card: {e}")
        return
    # fill links that the model may have missed
    entry.setdefault("links", {})
    entry["links"]["github"] = entry["links"].get("github") or (github_url or "")
    entry["links"]["article"] = entry["links"].get("article") or (article_url or "")
    entry["links"]["website"] = entry["links"].get("website") or (website or "")

    added, reason = add_resource(entry)
    await status.delete()
    if added:
        await update.message.reply_text(
            f"✅ Saved tool card: <b>{entry.get('title','')}</b>",
            parse_mode=ParseMode.HTML)
    else:
        await update.message.reply_text(
            f"⚠️ Skipped — {reason} (already in your knowledge base).")
    await update.message.reply_text(_pretty_json(entry))
    await update.message.reply_text(_pages_link(), parse_mode=ParseMode.HTML)


async def _handle_pdf(update, status, data):
    """PDF / Word doc: full descriptive summary, saved as a resource card.

    The user wants a complete, readable explanation they can revisit on the
    web page — so we generate the full deep-dive (not the short summary) and
    persist it as a resource.
    """
    kind = "Word document" if data.get("is_doc") else "PDF"
    if data.get("note"):  # e.g. legacy .doc not supported
        await status.edit_text(f"⚠️ {data['note']}")
        return
    await status.edit_text(f"📄 Reading the {kind} and writing a full explanation…")
    try:
        summary = synthesize(data, mode="full", num_predict=4096)
    except Exception as e:  # noqa: BLE001
        logger.exception("doc summary failed")
        await status.edit_text(f"❌ Couldn't summarize the {kind}: {e}")
        return
    if not summary:
        await status.edit_text(f"⚠️ Couldn't read that {kind}. Paste the text.")
        return

    entry = {
        "title": data.get("title") or (data.get("url") or kind),
        "description": summary,
        "type": "doc" if data.get("is_doc") else "paper",
        "links": {"github": "", "website": "",
                  "article": data.get("url") or ""},
    }
    added, reason = add_resource(entry)
    await status.delete()
    if added:
        await update.message.reply_text(
            f"✅ Saved full summary: <b>{entry['title']}</b>",
            parse_mode=ParseMode.HTML)
    else:
        await update.message.reply_text(
            f"⚠️ Skipped — {reason} (already in your knowledge base).")
    await _send_html(update, summary)
    await update.message.reply_text(_pages_link(), parse_mode=ParseMode.HTML)


async def _handle_learning(update, status, article_text, article_url):
    prompt = build_learning_json_prompt(article_text, article_url)
    try:
        entry = synthesize_json(prompt, num_predict=800)
    except Exception as e:  # noqa: BLE001
        logger.exception("learning json failed")
        await status.edit_text(f"❌ Couldn't build the learning card: {e}")
        return
    entry["links"] = entry.get("links") or article_url
    added, reason = add_learning(entry)
    await status.delete()
    if added:
        await update.message.reply_text("✅ Saved learning card.",
                                        parse_mode=ParseMode.HTML)
    else:
        await update.message.reply_text(
            f"⚠️ Skipped — {reason} (already in your knowledge base).")
    await update.message.reply_text(_pretty_json(entry))
    await update.message.reply_text(_pages_link(), parse_mode=ParseMode.HTML)


# ----------------------------------------------------------------- helpers
def _find_github_link(text):
    m = re.search(r"https?://github\.com/[\w.-]+/[\w.-]+", text or "")
    return m.group(0) if m else None


def _looks_like_learning(text):
    """Heuristic: upskilling / productivity / self-improvement article."""
    keys = ("productiv", "upskill", "improve", "habit", "career", "learn",
            "routine", "focus", "suggest", "tip", "advice", "growth",
            "well-being", "wellbeing", "mental", "burnout")
    t = (text or "").lower()
    return sum(k in t for k in keys) >= 2


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle a PDF / .docx file sent as a Telegram attachment."""
    doc = update.message.document
    fname = (doc.file_name or "").lower()
    status = await update.message.reply_text("📎 Downloading file…")
    try:
        tf = await doc.get_file()
        import tempfile, os
        suffix = ".pdf" if fname.endswith(".pdf") else \
            ".docx" if fname.endswith(".docx") else os.path.splitext(fname)[1] or ".bin"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
            path = f.name
        await tf.download_to_drive(path)
        data = {"source": "Article", "title": doc.file_name or "Uploaded file",
                "url": None, "text": None}
        if fname.endswith(".pdf"):
            from pypdf import PdfReader
            with open(path, "rb") as fh:
                reader = PdfReader(fh)
                text = "\n\n".join((p.extract_text() or "") for p in reader.pages).strip()
                title = (reader.metadata.title if reader.metadata else None) or doc.file_name
            data.update({"text": text[:30000], "is_pdf": True, "title": title})
        elif fname.endswith(".docx"):
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            data.update({"text": "\n\n".join(parts).strip()[:30000], "is_doc": True})
        else:
            await status.edit_text("⚠️ Only PDF and .docx files are supported. "
                                   "Re-send as one of those.")
            os.unlink(path)
            return
        os.unlink(path)
        if not data.get("text"):
            await status.edit_text("⚠️ Couldn't extract text from that file.")
            return
        await _handle_pdf(update, status, data)
    except Exception as e:  # noqa: BLE001
        logger.exception("document handling failed")
        await status.edit_text(f"❌ Couldn't read that file: {e}")


def main():
    if not TELEGRAM_BOT_TOKEN:
        print("ERROR: set TELEGRAM_BOT_TOKEN in your .env first.")
        return
    # Ensure the site data files reflect current storage on startup.
    try:
        from content_digest_bot.store import _regen
        _regen()
    except Exception as e:  # noqa: BLE001
        logger.warning("Could not regenerate site data: %s", e)
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_cmd))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle))
    app.add_handler(MessageHandler(filters.Document.PDF | filters.Document.FileExtension("docx"), handle_document))
    print(f"Bot running (model={ANTHROPIC_MODEL})…  Ctrl+C to stop.")
    # launchd restarts the process on crash; a manual retry loop here can spawn
    # overlapping pollers (double getUpdates → 409), so just run once.
    app.run_polling()


if __name__ == "__main__":
    main()
